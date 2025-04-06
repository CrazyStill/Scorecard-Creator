import os, csv, json, shutil, tempfile
from flask import Flask, render_template, request, redirect, url_for, flash, send_file, make_response
from werkzeug.utils import secure_filename
from docx import Document
import docx2pdf
from PyPDF2 import PdfMerger
import comtypes.client  # Helps with error handling via fallbacks
from io import BytesIO
import pythoncom       # Initializes COM


app = Flask(__name__)
app.secret_key = 'your_secret_key'  # Facilitates user messages

# Base directory for storing uploaded scorecard templates
BASE_DIR = os.path.abspath(os.path.dirname(__file__))
SCTEMP_DIR = os.path.join(BASE_DIR, 'SCTEMP')
if not os.path.exists(SCTEMP_DIR):
    os.makedirs(SCTEMP_DIR)

def allowed_file(filename, allowed_extensions):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions

# Functions for processing documents

def merge_runs_in_paragraph(paragraph):
    if len(paragraph.runs) <= 1:
        return
    full_text = "".join(run.text for run in paragraph.runs)
    for run in paragraph.runs:
        run.text = ""
    paragraph.runs[0].text = full_text

def replace_text_in_paragraph(paragraph, placeholder, replacement):
    if placeholder in paragraph.text:
        new_text = paragraph.text.replace(placeholder, replacement)
        paragraph.runs[0].text = new_text

def replace_placeholders_in_paragraphs(paragraphs, placeholders):
    for paragraph in paragraphs:
        merge_runs_in_paragraph(paragraph)
        for placeholder, replacement in placeholders:
            replace_text_in_paragraph(paragraph, placeholder, replacement)

# Replaces placeholders in CSV doc
def replace_text_in_doc(doc, rows, mapping, cards_per_page=4):
    all_placeholders = []
    for i in range(1, cards_per_page+1):
        row = rows[i-1] if i-1 < len(rows) else None
        for csv_header, placeholder in mapping.items():
            value = row.get(csv_header, "") if row else ""
            all_placeholders.append((f"{placeholder}_{i}", value))
    replace_placeholders_in_paragraphs(doc.paragraphs, all_placeholders)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholders_in_paragraphs(cell.paragraphs, all_placeholders)

def merge_two_pdfs(pdf1, pdf2, merged_pdf):
    merger = PdfMerger()
    merger.append(pdf1)
    merger.append(pdf2)
    merger.write(merged_pdf)
    merger.close()

def convert_docx_to_pdf(docx_path, pdf_path):
    # If conversion to PDF fails, fallback is used for error handling
    try:
        docx2pdf.convert(docx_path, pdf_path)
    except Exception as e:
        # COM fallback code
        pythoncom.CoInitialize()
        wdFormatPDF = 17
        try:
            word = comtypes.client.CreateObject("Word.Application")
        except Exception as e2:
            pythoncom.CoUninitialize()
            raise e2
        word.Visible = False
        doc = word.Documents.Open(docx_path)
        try:
            doc.SaveAs2(pdf_path, FileFormat=wdFormatPDF)
        except Exception as e3:
            doc.Close()
            word.Quit()
            pythoncom.CoUninitialize()
            raise e3
        doc.Close()
        word.Quit()
        pythoncom.CoUninitialize()

# Generates scorecard PDF, stores temporary files in temp_dir
def generate_scorecard(template_path, csv_path, mapping, cards_per_page=4, back_pdf_path=None, temp_dir=None):
    final_pdf_list = []
    with open(csv_path, newline='', encoding='latin-1') as f:
        sample = f.read(1024)
        f.seek(0)
        dialect = csv.Sniffer().sniff(sample)
        reader = csv.DictReader(f, dialect=dialect)
        rows = [row for row in reader if any(value.strip() for value in row.values())]
    for i in range(0, len(rows), cards_per_page):
        group = rows[i:i+cards_per_page]
        front_doc = Document(template_path)
        replace_text_in_doc(front_doc, group, mapping, cards_per_page=cards_per_page)
        temp_front_docx = os.path.join(temp_dir, f"temp_front_{i//cards_per_page}.docx")
        temp_front_pdf = os.path.join(temp_dir, f"temp_front_{i//cards_per_page}.pdf")
        page_pdf = os.path.join(temp_dir, f"page_{i//cards_per_page}.pdf")
        front_doc.save(temp_front_docx)
        convert_docx_to_pdf(temp_front_docx, temp_front_pdf)
        if back_pdf_path and os.path.exists(back_pdf_path):
            merge_two_pdfs(temp_front_pdf, back_pdf_path, page_pdf)
        else:
            shutil.copy(temp_front_pdf, page_pdf)
        final_pdf_list.append(page_pdf)
        os.remove(temp_front_docx)
        os.remove(temp_front_pdf)
    output_pdf = os.path.join(temp_dir, "Final_Scorecards.pdf")
    merger = PdfMerger()
    sorted_pdf_list = sorted(final_pdf_list, key=lambda x: int(x.split('_')[-1].split('.')[0]))
    for pdf in sorted_pdf_list:
        merger.append(pdf)
    merger.write(output_pdf)
    merger.close()
    for pdf in final_pdf_list:
        os.remove(pdf)
    return output_pdf

# Routes for Flask application

# Lists available templates
@app.route('/')
def index():
    templates_list = []
    for sport in os.listdir(SCTEMP_DIR):
        sport_dir = os.path.join(SCTEMP_DIR, sport)
        if os.path.isdir(sport_dir):
            for template_name in os.listdir(sport_dir):
                template_dir = os.path.join(sport_dir, template_name)
                if os.path.isdir(template_dir):
                    templates_list.append({'sport': sport, 'template': template_name})
    return render_template('index.html', templates=templates_list)

# Uploads new .docx template along with CSV (and optionally a PDF back page)
@app.route('/upload', methods=['GET', 'POST'])
def upload():
    if request.method == 'POST':
        sport = request.form.get('sport')
        template_name = request.form.get('template_name')
        if not sport or not template_name:
            flash("Sport and Template Name are required.")
            return redirect(request.url)
        sport = secure_filename(sport)
        template_name = secure_filename(template_name)
        template_dir = os.path.join(SCTEMP_DIR, sport, template_name)
        os.makedirs(template_dir, exist_ok=True)
        # Save .docx template for front
        front_file = request.files.get('front_file')
        if not front_file or not allowed_file(front_file.filename, {'docx'}):
            flash("A valid Word template file (.docx) is required.")
            return redirect(request.url)
        front_path = os.path.join(template_dir, 'template_front.docx')
        front_file.save(front_path)
        # Save CSV for data entry
        csv_file = request.files.get('csv_file')
        if not csv_file or not allowed_file(csv_file.filename, {'csv'}):
            flash("A valid CSV file is required.")
            return redirect(request.url)
        csv_path = os.path.join(template_dir, 'template_data.csv')
        csv_file.save(csv_path)
        # Optional PDF for back
        back_option = request.form.get('back_option')
        if back_option == 'yes':
            back_file = request.files.get('back_file')
            if back_file and allowed_file(back_file.filename, {'pdf'}):
                back_path = os.path.join(template_dir, 'template_back.pdf')
                back_file.save(back_path)
            else:
                flash("Back design selected but no valid PDF uploaded.")
                return redirect(request.url)
        return redirect(url_for('mapping', sport=sport, template_name=template_name))
    return render_template('upload.html')

# Reads CSV headers and maps them to placeholder text, then saves mapping as JSON
@app.route('/mapping/<sport>/<template_name>', methods=['GET', 'POST'])
def mapping(sport, template_name):
    template_dir = os.path.join(SCTEMP_DIR, secure_filename(sport), secure_filename(template_name))
    csv_path = os.path.join(template_dir, 'template_data.csv')
    if not os.path.exists(csv_path):
        flash("CSV file not found for this template.")
        return redirect(url_for('index'))
    with open(csv_path, newline='', encoding='latin-1') as f:
        sample = f.read(1024)
        f.seek(0)
        dialect = csv.Sniffer().sniff(sample)
        reader = csv.reader(f, dialect=dialect)
        headers = next(reader)
    if request.method == 'POST':
        mapping_data = {}
        for header in headers:
            placeholder = request.form.get(header)
            mapping_data[header] = placeholder if placeholder else header
        try:
            cards_per_page = int(request.form.get('cards_per_page', 4))
        except ValueError:
            cards_per_page = 4
        mapping_file = os.path.join(template_dir, 'mapping.json')
        with open(mapping_file, 'w') as f:
            json.dump({"cards_per_page": cards_per_page, "mapping": mapping_data}, f)
        flash("Mapping saved successfully.")
        return redirect(url_for('index'))
    instructions = ("For each phrase to be replaced on the template, append an underscore and a number "
                    "corresponding to the scorecard position on the page (e.g., if 3 scorecards per page, use _1, _2, _3).")
    return render_template('mapping.html', headers=headers, sport=sport, template_name=template_name, instructions=instructions)

# About page route
@app.route('/about')
def about():
    return render_template('about.html')

# New PDF preview route: Converts DOCX to PDF (merging optional back PDF if available) and returns the PDF inline
@app.route('/preview_pdf/<sport>/<template_name>')
def preview_pdf(sport, template_name):
    template_dir = os.path.join(SCTEMP_DIR, secure_filename(sport), secure_filename(template_name))
    front_docx = os.path.join(template_dir, 'template_front.docx')
    if not os.path.exists(front_docx):
        flash("DOCX template not found.")
        return redirect(url_for('index'))

    # Create a temporary directory for conversion
    temp_dir = tempfile.mkdtemp()
    temp_front_pdf = os.path.join(temp_dir, "temp_front.pdf")
    try:
        convert_docx_to_pdf(front_docx, temp_front_pdf)
    except Exception as e:
        shutil.rmtree(temp_dir)
        return f"Error converting DOCX to PDF: {e}", 500

    back_pdf = os.path.join(template_dir, 'template_back.pdf')
    if os.path.exists(back_pdf):
        preview_pdf_path = os.path.join(temp_dir, "preview.pdf")
        merge_two_pdfs(temp_front_pdf, back_pdf, preview_pdf_path)
        pdf_to_send = preview_pdf_path
    else:
        pdf_to_send = temp_front_pdf

    # Read the PDF into memory before cleaning up
    with open(pdf_to_send, 'rb') as f:
        pdf_bytes = f.read()
    shutil.rmtree(temp_dir)
    return send_file(BytesIO(pdf_bytes),
                     mimetype='application/pdf',
                     download_name='preview.pdf',
                     as_attachment=False)

# Modified preview route: Renders a page with an embedded PDF viewer for previewing the template
@app.route('/preview/<sport>/<template_name>', methods=['GET', 'POST'])
def preview(sport, template_name):
    template_dir = os.path.join(SCTEMP_DIR, secure_filename(sport), secure_filename(template_name))
    docx_path = os.path.join(template_dir, 'template_front.docx')
    if request.method == 'POST':
        new_docx = request.files.get('new_docx')
        if new_docx and allowed_file(new_docx.filename, {'docx'}):
            new_docx.save(docx_path)
            flash("Template updated successfully.")
            return redirect(url_for('preview', sport=sport, template_name=template_name))
        else:
            flash("Please upload a valid DOCX file.")
            return redirect(url_for('preview', sport=sport, template_name=template_name))
    return render_template('preview.html', sport=sport, template_name=template_name)

# Downloads the current DOCX template
@app.route('/download_template/<sport>/<template_name>')
def download_template(sport, template_name):
    template_dir = os.path.join(SCTEMP_DIR, secure_filename(sport), secure_filename(template_name))
    docx_path = os.path.join(template_dir, 'template_front.docx')
    if not os.path.exists(docx_path):
         flash("DOCX template not found.")
         return redirect(url_for('index'))
    return send_file(docx_path, as_attachment=True)

# Generates the final scorecard PDF using the uploaded CSV data and mapping; sets a cookie for redirection
@app.route('/generate/<sport>/<template_name>', methods=['GET', 'POST'])
def generate(sport, template_name):
    template_dir = os.path.join(SCTEMP_DIR, secure_filename(sport), secure_filename(template_name))
    front_path = os.path.join(template_dir, 'template_front.docx')
    mapping_file = os.path.join(template_dir, 'mapping.json')
    back_path = os.path.join(template_dir, 'template_back.pdf')
    if request.method == 'POST':
        filled_csv_file = request.files.get('filled_csv')
        if not filled_csv_file or not allowed_file(filled_csv_file.filename, {'csv'}):
            flash("Please upload a valid CSV file with your data.")
            return redirect(request.url)
        with tempfile.TemporaryDirectory() as temp_dir:
            filled_csv_path = os.path.join(temp_dir, 'filled_data.csv')
            filled_csv_file.save(filled_csv_path)
            with open(mapping_file, 'r') as f:
                mapping_json = json.load(f)
            cards_per_page = mapping_json.get("cards_per_page", 4)
            mapping_data = mapping_json.get("mapping", {})
            output_pdf = generate_scorecard(front_path, filled_csv_path, mapping_data,
                                            cards_per_page=cards_per_page,
                                            back_pdf_path=back_path if os.path.exists(back_path) else None,
                                            temp_dir=temp_dir)
            with open(output_pdf, 'rb') as pdf_file:
                pdf_bytes = pdf_file.read()
            response = make_response(send_file(BytesIO(pdf_bytes),
                             download_name=os.path.basename(output_pdf),
                             mimetype='application/pdf',
                             as_attachment=True))
            response.set_cookie("fileDownload", "true", max_age=60)
            return response
    else:
        return render_template('generate.html', sport=sport, template_name=template_name)

# Shows the original CSV template for download
@app.route('/download_csv/<sport>/<template_name>')
def download_csv(sport, template_name):
    template_dir = os.path.join(SCTEMP_DIR, secure_filename(sport), secure_filename(template_name))
    csv_template_path = os.path.join(template_dir, 'template_data.csv')
    if not os.path.exists(csv_template_path):
         flash("CSV template not found.")
         return redirect(url_for('index'))
    return send_file(csv_template_path, as_attachment=True)

# Completely wipes the given template directory
@app.route('/delete/<sport>/<template_name>', methods=['POST'])
def delete_template(sport, template_name):
    template_dir = os.path.join(SCTEMP_DIR, secure_filename(sport), secure_filename(template_name))
    if os.path.exists(template_dir):
        shutil.rmtree(template_dir)
        flash("Template deleted successfully.")
    else:
        flash("Template not found.")
    return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(debug=True)
